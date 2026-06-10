# -*- coding: utf-8 -*-
"""
03_generate_word.py
根据 01_analysis.py 输出的 results.json 和图片，生成正式 Word 作业文档。

运行前请先运行：
    python 01_analysis.py

然后运行：
    python 03_generate_word.py
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = PROJECT_ROOT / "outputs"
REPORT_DIR = PROJECT_ROOT / "reports"
RESULTS_PATH = OUTPUT_DIR / "results.json"
WORD_PATH = REPORT_DIR / "管理统计学_广告战略方差分析作业.docx"

COLORS = {
    "navy": "0F2742",
    "blue": "2D6CDF",
    "sky": "5AA9E6",
    "light_blue": "EAF4FF",
    "header": "D9EAF7",
    "answer_bg": "F3F8FF",
    "green_bg": "F0FDF4",
    "orange_bg": "FFF7E6",
}


def load_results() -> Dict[str, Any]:
    if not RESULTS_PATH.exists():
        raise FileNotFoundError("找不到 outputs/results.json。请先运行 python 01_analysis.py")
    with open(RESULTS_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def hex_to_rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.replace("#", "")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_run_font(
    run,
    font_name: str = "宋体",
    size: int | None = None,
    bold: bool | None = None,
    color: str | None = None,
) -> None:
    """设置中英文兼容字体。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = hex_to_rgb(color)


def set_paragraph_text(paragraph, text: str, font_name: str = "宋体", size: int = 11, bold: bool = False, color: str | None = None) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, font_name=font_name, size=size, bold=bold, color=color)


def set_cell_shading(cell, fill: str) -> None:
    """设置单元格底色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    """写入单元格文本并设置字体。"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def fmt(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        return f"{value:.4f}"
    return str(value)


def num(value: Any, ndigits: int = 4) -> str:
    if value is None:
        return ""
    try:
        return f"{float(value):.{ndigits}f}"
    except (TypeError, ValueError):
        return str(value)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if not p.runs:
        run = p.add_run(text)
    else:
        run = p.runs[0]
    set_run_font(run, font_name="黑体", size=16 if level == 1 else 13, bold=True, color=COLORS["navy"])


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.25
    set_paragraph_text(p, text, size=11)


def add_plain_paragraph(doc: Document, text: str, bold: bool = False, color: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    set_paragraph_text(p, text, size=11, bold=bold, color=color)


def add_answer_box(doc: Document, title: str, body: str, direct_answer: str) -> None:
    """添加逐题作答框，让老师一眼看到答案。"""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, COLORS["answer_bg"])
    cell.text = ""

    p_title = cell.paragraphs[0]
    r_title = p_title.add_run(title)
    set_run_font(r_title, font_name="黑体", size=11, bold=True, color=COLORS["blue"])

    p_body = cell.add_paragraph()
    p_body.paragraph_format.line_spacing = 1.2
    r_body = p_body.add_run(body)
    set_run_font(r_body, font_name="宋体", size=10)

    p_answer = cell.add_paragraph()
    p_answer.paragraph_format.line_spacing = 1.2
    r_answer = p_answer.add_run("直接答案：" + direct_answer)
    set_run_font(r_answer, font_name="宋体", size=10, bold=True, color=COLORS["navy"])

    doc.add_paragraph()


def add_direct_conclusion_box(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, COLORS["green_bg"])
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run, font_name="宋体", size=10, bold=True, color=COLORS["navy"])
    doc.add_paragraph()


def prevent_row_split(row) -> None:
    """尽量避免 Word 表格行跨页拆分。"""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def select_records(records: List[Dict[str, Any]], columns: List[tuple[str, str]]) -> List[Dict[str, Any]]:
    """按需要选择并重命名表格列，避免 Word 中表格过宽。"""
    selected = []
    for record in records:
        selected.append({label: record.get(key) for key, label in columns})
    return selected


def add_records_table(doc: Document, records: List[Dict[str, Any]], title: str, font_size: int = 9) -> None:
    """添加数据表。"""
    add_plain_paragraph(doc, title, bold=True, color=COLORS["navy"])
    if not records:
        add_body_paragraph(doc, "无数据。")
        return

    columns = list(records[0].keys())
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    prevent_row_split(table.rows[0])
    for i, col in enumerate(columns):
        set_cell_text(header_cells[i], str(col), bold=True, size=font_size)
        set_cell_shading(header_cells[i], COLORS["header"])

    for record in records:
        row = table.add_row()
        prevent_row_split(row)
        row_cells = row.cells
        for i, col in enumerate(columns):
            set_cell_text(row_cells[i], fmt(record.get(col)), size=font_size)

    doc.add_paragraph()


def add_picture(doc: Document, relative_path: str, title: str, width_cm: float = 14.5) -> None:
    """添加图片及题注。"""
    path = PROJECT_ROOT / relative_path
    if not path.exists():
        raise FileNotFoundError(f"找不到图片：{path}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(cap, title, size=10)


def setup_document_style(doc: Document) -> None:
    """设置页面和默认样式。"""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)


def build_word(results: Dict[str, Any]) -> Document:
    doc = Document()
    setup_document_style(doc)

    s624 = results["summary_624"]
    s625 = results["summary_625"]
    interp = results["interpretation"]
    alpha = results.get("alpha", 0.05)

    f624 = num(s624.get("f_stat"))
    p624 = num(s624.get("p_value"))
    p_strategy = num(s625.get("strategy_p"))
    p_media = num(s625.get("media_p"))
    p_interaction = num(s625.get("interaction_p"))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("管理统计学作业：广告战略与广告媒体的方差分析")
    set_run_font(run, font_name="黑体", size=18, bold=True, color=COLORS["navy"])

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(subtitle, "基于表6-24与表6-25的销售量数据", size=11)

    add_heading(doc, "一、讨论题直接作答", 1)
    add_answer_box(
        doc,
        "讨论题1：根据案例给出显著性检验假设",
        "针对表6-24，原假设H0为 μ便利性 = μ高质量 = μ低价格，备择假设H1为三种广告战略下的平均周销量不全相同。针对表6-25，应分别检验广告策略主效应、广告媒体主效应和广告策略×广告媒体交互作用。",
        "本题应设置单因素方差分析假设和双因素方差分析假设。"
    )
    add_answer_box(
        doc,
        "讨论题2：检验三种不同广告战略下的销售量均值是否相同",
        f"单因素方差分析结果为 F = {f624}，p = {p624}。由于 p = {p624} > {alpha}，在5%显著性水平下不能拒绝原假设。",
        "三种广告战略下的销售量均值没有显著差异。"
    )
    add_answer_box(
        doc,
        "讨论题3：检验三种广告策略和两种广告媒体组合下的销售量均值是否相同",
        f"双因素方差分析结果为：广告策略主效应 p = {p_strategy}，广告媒体主效应 p = {p_media}，策略×媒体交互作用 p = {p_interaction}。三者均大于0.05。",
        "在5%显著性水平下，广告策略、广告媒体及其交互作用均不显著，不能认为各组合销售量均值存在显著差异。"
    )
    add_answer_box(
        doc,
        "讨论题4：对上述检验方法进行讨论",
        "表6-24只有广告战略一个因素，适合采用单因素方差分析；表6-25包含广告策略和广告媒体两个因素，适合采用双因素方差分析。但方差分析依赖独立性、近似正态性和方差齐性，现实中的城市差异和周销量时间相关性可能影响检验结果。",
        "方差分析方法适用，但检验结论应结合样本量、城市控制和时间相关性谨慎解释。"
    )

    add_heading(doc, "二、问题背景与研究目的", 1)
    add_body_paragraph(
        doc,
        "某厂家开发了一种具有便利性、高质量和低价格优势的新产品。营销经理需要判断应当采用何种广告战略，并进一步考察不同广告媒体下广告策略的销售效果。本作业基于题目给出的周销量数据，采用方差分析方法检验不同广告方案下销售量均值是否存在显著差异。"
    )

    add_heading(doc, "三、变量说明与试验设计", 1)
    add_body_paragraph(
        doc,
        "表6-24包含三种广告战略下20个星期的周销量数据，因素为广告战略，水平包括便利性、高质量和低价格，适合采用单因素方差分析。表6-25包含三种广告策略和两种广告媒体组合下10个星期的周销量数据，因素包括广告策略和广告媒体，适合采用双因素方差分析，并检验二者是否存在交互作用。"
    )

    add_heading(doc, "四、显著性检验假设展开", 1)
    add_plain_paragraph(doc, "1. 单因素方差分析假设", bold=True, color=COLORS["navy"])
    add_body_paragraph(doc, "原假设H0：三种广告战略下的平均周销量相同，即 μ便利性 = μ高质量 = μ低价格。")
    add_body_paragraph(doc, "备择假设H1：三种广告战略下的平均周销量不全相同。")

    add_plain_paragraph(doc, "2. 双因素方差分析假设", bold=True, color=COLORS["navy"])
    add_body_paragraph(doc, "广告策略主效应：原假设为三种广告策略下的平均周销量相同；备择假设为三者不全相同。")
    add_body_paragraph(doc, "广告媒体主效应：原假设为两种广告媒体下的平均周销量相同；备择假设为两者不同。")
    add_body_paragraph(doc, "交互作用：原假设为广告策略与广告媒体之间不存在交互作用；备择假设为二者存在交互作用。")

    doc.add_page_break()
    add_heading(doc, "五、三种广告战略销量均值比较：单因素方差分析", 1)
    add_records_table(doc, select_records(results["desc_624"], [("strategy", "广告战略"), ("n", "n"), ("mean", "均值"), ("std", "标准差"), ("min", "最小值"), ("max", "最大值")]), "表1  表6-24描述性统计")
    add_picture(doc, results["figures"]["fig_1_strategy_mean"], "图1  三种广告战略的平均周销量及95%置信区间")
    add_picture(doc, results["figures"]["fig_2_strategy_boxplot"], "图2  三种广告战略的周销量分布箱线图")
    add_records_table(doc, results["anova_624"], "表2  单因素方差分析表")
    add_direct_conclusion_box(doc, f"直接结论：F = {f624}，p = {p624}。在 α = {alpha} 下，p值大于0.05，因此三种广告战略均值不存在显著差异。")
    add_body_paragraph(doc, f"辅助检验结果显示，方差齐性检验Levene p = {num(s624.get('levene_p'))}，残差正态性检验Shapiro-Wilk p = {num(s624.get('shapiro_p'))}。这些结果用于辅助判断方差分析前提是否基本满足。")

    doc.add_page_break()
    add_heading(doc, "六、广告策略与广告媒体组合效果分析：双因素方差分析", 1)
    add_records_table(doc, select_records(results["desc_625"], [("strategy", "广告策略"), ("media", "广告媒体"), ("n", "n"), ("mean", "均值"), ("std", "标准差"), ("min", "最小值"), ("max", "最大值")]), "表3  表6-25各组合描述性统计")
    add_picture(doc, results["figures"]["fig_3_combination_mean"], "图3  广告策略×广告媒体组合的平均周销量")
    add_picture(doc, results["figures"]["fig_4_interaction"], "图4  广告策略与广告媒体组合的平均周销量对比")
    add_records_table(doc, results["anova_625"], "表4  双因素方差分析表")
    add_direct_conclusion_box(doc, f"直接结论：广告策略 p = {p_strategy}，广告媒体 p = {p_media}，策略×媒体交互作用 p = {p_interaction}。三者均大于0.05，因此广告策略、广告媒体和二者交互作用均不显著。")
    add_body_paragraph(doc, f"广告媒体的p值接近0.05，说明社交网站可能有更高销量表现，但在5%显著性水平下证据仍不足。辅助检验结果显示，方差齐性检验Levene p = {num(s625.get('levene_p'))}，残差正态性检验Shapiro-Wilk p = {num(s625.get('shapiro_p'))}。")

    add_heading(doc, "七、管理解释与建议", 1)
    add_body_paragraph(doc, interp["management"])

    add_heading(doc, "八、检验方法讨论", 1)
    add_body_paragraph(doc, interp["method_discussion"])

    add_heading(doc, "九、最终结论", 1)
    add_body_paragraph(
        doc,
        "综合单因素方差分析和双因素方差分析结果，在5%显著性水平下，本题数据尚不足以证明三种广告战略、两种广告媒体以及二者交互作用会导致平均周销量的显著差异。描述性统计结果可以提示可能的市场方向，但正式营销决策仍需要更大样本、更长观察期和更严格的试验控制。"
    )

    return doc


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    results = load_results()
    doc = build_word(results)
    doc.save(WORD_PATH)
    print(f"Word文档已生成：{WORD_PATH}")


if __name__ == "__main__":
    main()
